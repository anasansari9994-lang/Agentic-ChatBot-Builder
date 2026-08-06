import uuid
import os
import pandas as pd
from datetime import datetime
from core.logging import logger
from docx import Document as DocxReader
from document_schema import Document, Elements, FileTypeEnum, ElementTypeEnum
from models.analysis_funciotn.image_analyser import describe_chart
from models.analysis_funciotn.table_analyser import generate_table_description
from parsers.base import BaseParser
from pathlib import Path

class DOCXParser(BaseParser):
    async def parse(self, file_path: str):
        logger.info("Docx Parser is running")
        doc = DocxReader(file_path)
        elemnets = []
        document_id =  str(uuid.uuid4())
        image_idx = 0
        logger.info("now start the docx parser")
        for para in doc.paragraphs:
            if para.text.strip():
                is_heading = para.style.name.startswith('Heading')
                logger.info(f"extract the element {is_heading}")
                elemnets.append(
                    Elements(
                        element_id=str(uuid.uuid4()),
                        element_type=ElementTypeEnum.TEXT,
                        content = para.text.strip(),
                        metadata = {
                            "is_heading" : is_heading,
                            "style" : para.style.name
                        }
                    )
                )
        for table_idx, table in enumerate(doc.tables):
            if table_idx == 0:
                logger.info("Table extraction is runing")
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                table_df = pd.DataFrame(table_data)
                table_md = table_df.to_markdown(index=False)
                table_summary = generate_table_description(table_md)
                elemnets.append(
                    Elements(
                        element_id=str(uuid.uuid4()),
                        element_type=ElementTypeEnum.TABLE,
                        content=table_summary,
                        metadata = {
                            "table_index" : table_idx,
                            "total_rows" : len(table.rows),
                            "total_cols" : len(table.columns)
                        }
                    )
                )

            # image_idx = 0
        logged = False
        image_dir = Path("storage/images") / document_id
        image_dir.mkdir(parents=True, exist_ok=True)
        for image_idx, shape in enumerate(doc.inline_shapes, start=1):
            if not logged:
                logger.info("image extraction is running")
                logged = True
            if shape.type != WD_INLINE_SHAPE.PICTURE:
                continue
            image_path = None
            try:
                r_id = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
                image_part = doc.part.related_parts[r_id]
                image_byte = image_part.blob
                extention = image_part.content_type.split("/")[-1]
                image_path = image_dir /f"image_{image_idx}.{extention}"
                with open(image_path, "wb") as f:
                    f.write(image_byte)
                description = describe_chart(str(image_path))
                content = description
            except Exception as e:
                logger.exception(f"error in processing image {image_idx}")
                content = ""

            elemnets.append(
                Elements(
                    element_id = str(uuid.uuid4()),
                    element_type = ElementTypeEnum.IMAGE,
                    content= content,
                    metadata={
                        "image_index": image_idx,
                        "image_path" : str(image_path)
                    }
                )
            )

        return Document(
            document_id = document_id,
            filename = os.path.basename(file_path),
            file_type = FileTypeEnum.DOCX,
            created_at = datetime.now(),
            metadata={
                "total_paragraphs": len(doc.paragraphs),
                "total_tables": len(doc.tables),
                "total_images": image_idx
            },
            elements=elemnets
        )


    